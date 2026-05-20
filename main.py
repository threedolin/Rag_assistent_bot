import os
import time
import hashlib
from pathlib import Path
from typing import List, Tuple

import streamlit as st
from dotenv import load_dotenv

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_chroma import Chroma

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_ollama import ChatOllama, OllamaEmbeddings


# =========================
# Базовая настройка проекта
# =========================

load_dotenv()

APP_TITLE = "RAG-ассистент по документам"
PERSIST_DIR = "./chroma_db"

st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📄",
    layout="wide",
)


# =========================
# Вспомогательные функции
# =========================

def init_session_state() -> None:
    """Инициализация переменных текущей сессии Streamlit."""
    if "messages" not in st.session_state:
        st.session_state.messages = []

    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None

    if "processed_info" not in st.session_state:
        st.session_state.processed_info = None


def read_txt(uploaded_file) -> str:
    """Чтение TXT с попыткой нескольких кодировок."""
    raw = uploaded_file.getvalue()

    for encoding in ["utf-8", "cp1251", "latin-1"]:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    return raw.decode("utf-8", errors="ignore")


def load_pdf(uploaded_file) -> List[Document]:
    """Извлекает текст из PDF постранично."""
    docs = []

    uploaded_file.seek(0)
    reader = PdfReader(uploaded_file)

    for page_idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": uploaded_file.name,
                        "page": page_idx,
                        "file_type": "pdf",
                    },
                )
            )

    return docs


def load_docx(uploaded_file) -> List[Document]:
    """Извлекает текст из DOCX."""
    uploaded_file.seek(0)
    docx = DocxDocument(uploaded_file)

    paragraphs = []
    for p in docx.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": uploaded_file.name,
                "page": None,
                "file_type": "docx",
            },
        )
    ]


def load_txt(uploaded_file) -> List[Document]:
    """Извлекает текст из TXT."""
    text = read_txt(uploaded_file).strip()

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": uploaded_file.name,
                "page": None,
                "file_type": "txt",
            },
        )
    ]


def load_uploaded_files(uploaded_files) -> List[Document]:
    """Загружает PDF/DOCX/TXT и превращает их в LangChain Document."""
    all_docs = []

    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name.lower()

        if file_name.endswith(".pdf"):
            docs = load_pdf(uploaded_file)
        elif file_name.endswith(".docx"):
            docs = load_docx(uploaded_file)
        elif file_name.endswith(".txt"):
            docs = load_txt(uploaded_file)
        else:
            st.warning(f"Файл {uploaded_file.name} пропущен: неподдерживаемый формат.")
            docs = []

        all_docs.extend(docs)

    return all_docs


def split_documents(
    docs: List[Document],
    chunk_size: int,
    chunk_overlap: int,
) -> List[Document]:
    """Разбивает документы на чанки."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )

    chunks = splitter.split_documents(docs)

    for idx, chunk in enumerate(chunks, start=1):
        chunk.metadata["chunk_id"] = idx

    return chunks


def make_collection_name(uploaded_files, provider: str, embedding_model: str) -> str:
    """Создаёт уникальное имя коллекции Chroma для набора файлов."""
    hasher = hashlib.sha256()

    for uploaded_file in uploaded_files:
        hasher.update(uploaded_file.name.encode("utf-8"))
        hasher.update(str(uploaded_file.size).encode("utf-8"))
        hasher.update(uploaded_file.getvalue())

    hasher.update(provider.encode("utf-8"))
    hasher.update(embedding_model.encode("utf-8"))

    digest = hasher.hexdigest()[:16]
    timestamp = int(time.time())

    return f"rag_{digest}_{timestamp}"


def get_embeddings(
    provider: str,
    openai_api_key: str,
    openai_embedding_model: str,
    ollama_embedding_model: str,
    ollama_base_url: str,
):
    """Возвращает embedding-модель для OpenAI или Ollama."""
    if provider == "OpenAI":
        if openai_api_key:
            os.environ["OPENAI_API_KEY"] = openai_api_key

        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("Не найден OPENAI_API_KEY. Укажи ключ в интерфейсе или в .env.")

        return OpenAIEmbeddings(model=openai_embedding_model)

    return OllamaEmbeddings(
        model=ollama_embedding_model,
        base_url=ollama_base_url,
    )


def get_llm(
    provider: str,
    openai_api_key: str,
    openai_chat_model: str,
    ollama_chat_model: str,
    ollama_base_url: str,
    temperature: float = 0.1,
):
    """Возвращает LLM для генерации ответа."""
    if provider == "OpenAI":
        if openai_api_key:
            os.environ["OPENAI_API_KEY"] = openai_api_key

        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("Не найден OPENAI_API_KEY. Укажи ключ в интерфейсе или в .env.")

        return ChatOpenAI(
            model=openai_chat_model,
            temperature=temperature,
        )

    return ChatOllama(
        model=ollama_chat_model,
        base_url=ollama_base_url,
        temperature=temperature,
    )


def build_vectorstore(
    chunks: List[Document],
    embeddings,
    collection_name: str,
) -> Chroma:
    """Создаёт и сохраняет векторное хранилище Chroma."""
    Path(PERSIST_DIR).mkdir(parents=True, exist_ok=True)

    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=collection_name,
        persist_directory=PERSIST_DIR,
    )

    return vectorstore


def format_docs_for_prompt(docs: List[Document]) -> str:
    """Формирует контекст для LLM из найденных чанков."""
    formatted = []

    for idx, doc in enumerate(docs, start=1):
        source = doc.metadata.get("source", "unknown")
        page = doc.metadata.get("page")
        chunk_id = doc.metadata.get("chunk_id", "?")

        page_text = f", page {page}" if page else ""
        header = f"[Источник {idx}: {source}{page_text}, chunk {chunk_id}]"

        formatted.append(f"{header}\n{doc.page_content}")

    return "\n\n---\n\n".join(formatted)


def format_chat_history(max_messages: int = 6) -> str:
    """Берёт последние сообщения из истории, чтобы LLM понимала контекст диалога."""
    recent = st.session_state.messages[-max_messages:]

    lines = []
    for msg in recent:
        role = "Пользователь" if msg["role"] == "user" else "Ассистент"
        content = msg["content"]

        # Не перегружаем промпт слишком длинной историей
        if len(content) > 1000:
            content = content[:1000] + "..."

        lines.append(f"{role}: {content}")

    return "\n".join(lines)


def answer_with_rag(
    question: str,
    vectorstore: Chroma,
    llm,
    top_k: int,
) -> Tuple[str, List[Document]]:
    """Ищет релевантные чанки и генерирует ответ с источниками."""
    retrieved_docs = vectorstore.similarity_search(question, k=top_k)

    context = format_docs_for_prompt(retrieved_docs)
    history = format_chat_history()

    prompt = ChatPromptTemplate.from_template(
        """
Ты — аккуратный RAG-ассистент по документам.

Твоя задача:
1. Отвечать только на основе контекста из документов.
2. Если в документах нет ответа, честно сказать: "В загруженных документах я не нашёл точного ответа".
3. Не выдумывать факты.
4. По возможности указывать источники в формате [Источник 1], [Источник 2].
5. Отвечать на русском языке, если пользователь пишет по-русски.

История текущего диалога:
{history}

Контекст из документов:
{context}

Вопрос пользователя:
{question}

Ответ:
"""
    )

    chain = prompt | llm

    response = chain.invoke(
        {
            "history": history,
            "context": context,
            "question": question,
        }
    )

    return response.content, retrieved_docs


def answer_without_rag(question: str, llm) -> str:
    """Обычный ответ LLM без использования загруженных документов."""
    prompt = ChatPromptTemplate.from_template(
        """
Ответь на вопрос пользователя без использования загруженных документов.
Если вопрос зависит от конкретного документа, предупреди, что без RAG ответ может быть неточным.

Вопрос:
{question}

Ответ:
"""
    )

    chain = prompt | llm
    response = chain.invoke({"question": question})

    return response.content


def render_sources(docs: List[Document]) -> None:
    """Показывает найденные фрагменты документов в интерфейсе."""
    if not docs:
        return

    with st.expander("Показать найденные фрагменты документов"):
        for idx, doc in enumerate(docs, start=1):
            source = doc.metadata.get("source", "unknown")
            page = doc.metadata.get("page")
            chunk_id = doc.metadata.get("chunk_id", "?")

            page_text = f", страница {page}" if page else ""

            st.markdown(f"**Источник {idx}: {source}{page_text}, chunk {chunk_id}**")
            st.write(doc.page_content)
            st.divider()


def reset_chat() -> None:
    """Очищает историю чата текущей сессии."""
    st.session_state.messages = []


# =========================
# Интерфейс Streamlit
# =========================

init_session_state()

st.title("📄 RAG-ассистент по документам")
st.caption("Загрузи PDF/DOCX/TXT и задавай вопросы по содержимому документов.")

with st.sidebar:
    st.header("Настройки")

    provider = st.radio(
        "Провайдер модели",
        ["OpenAI", "Ollama"],
        help="OpenAI — облачный API. Ollama — локальный запуск на твоём компьютере.",
    )

    st.subheader("Модели")

    if provider == "OpenAI":
        openai_api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            value=os.getenv("OPENAI_API_KEY", ""),
            help="Можно не вводить, если ключ уже есть в переменной окружения OPENAI_API_KEY.",
        )

        openai_chat_model = st.text_input(
            "Chat model",
            value="gpt-4o-mini",
        )

        openai_embedding_model = st.text_input(
            "Embedding model",
            value="text-embedding-3-small",
        )

        ollama_chat_model = "llama3.2:3b"
        ollama_embedding_model = "nomic-embed-text"
        ollama_base_url = "http://localhost:11434"

    else:
        openai_api_key = ""
        openai_chat_model = "gpt-4o-mini"
        openai_embedding_model = "text-embedding-3-small"

        ollama_base_url = st.text_input(
            "Ollama base URL",
            value="http://localhost:11434",
        )

        ollama_chat_model = st.text_input(
            "Ollama chat model",
            value="llama3.2:3b",
        )

        ollama_embedding_model = st.text_input(
            "Ollama embedding model",
            value="nomic-embed-text",
        )

    st.subheader("Параметры RAG")

    chunk_size = st.slider(
        "Размер чанка",
        min_value=300,
        max_value=2000,
        value=900,
        step=100,
    )

    chunk_overlap = st.slider(
        "Перекрытие чанков",
        min_value=0,
        max_value=500,
        value=150,
        step=50,
    )

    top_k = st.slider(
        "Количество источников для поиска",
        min_value=1,
        max_value=8,
        value=4,
        step=1,
    )

    compare_without_rag = st.checkbox(
        "Фишка: сравнить ответ с RAG и без RAG",
        value=False,
    )

    st.divider()

    if st.button("Очистить историю чата"):
        reset_chat()
        st.rerun()


uploaded_files = st.file_uploader(
    "Загрузи документы",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True,
)

process_clicked = st.button("Обработать документы", type="primary")

if process_clicked:
    if not uploaded_files:
        st.error("Сначала загрузи хотя бы один PDF/DOCX/TXT файл.")
    else:
        try:
            with st.spinner("Читаю документы..."):
                raw_docs = load_uploaded_files(uploaded_files)

            if not raw_docs:
                st.error("Не удалось извлечь текст из документов.")
            else:
                with st.spinner("Разбиваю документы на чанки..."):
                    chunks = split_documents(raw_docs, chunk_size, chunk_overlap)

                embedding_model_name = (
                    openai_embedding_model if provider == "OpenAI" else ollama_embedding_model
                )

                with st.spinner("Создаю embeddings и сохраняю Chroma DB..."):
                    embeddings = get_embeddings(
                        provider=provider,
                        openai_api_key=openai_api_key,
                        openai_embedding_model=openai_embedding_model,
                        ollama_embedding_model=ollama_embedding_model,
                        ollama_base_url=ollama_base_url,
                    )

                    collection_name = make_collection_name(
                        uploaded_files,
                        provider,
                        embedding_model_name,
                    )

                    vectorstore = build_vectorstore(
                        chunks=chunks,
                        embeddings=embeddings,
                        collection_name=collection_name,
                    )

                st.session_state.vectorstore = vectorstore
                st.session_state.processed_info = {
                    "files": [file.name for file in uploaded_files],
                    "chunks": len(chunks),
                    "collection_name": collection_name,
                    "provider": provider,
                    "embedding_model": embedding_model_name,
                }

                st.success(
                    f"Готово! Обработано файлов: {len(uploaded_files)}, "
                    f"создано чанков: {len(chunks)}."
                )

        except Exception as e:
            st.error(f"Ошибка при обработке документов: {e}")


if st.session_state.processed_info:
    info = st.session_state.processed_info

    with st.expander("Информация о текущем индексе", expanded=False):
        st.write(f"**Файлы:** {', '.join(info['files'])}")
        st.write(f"**Чанков:** {info['chunks']}")
        st.write(f"**Провайдер:** {info['provider']}")
        st.write(f"**Embedding model:** {info['embedding_model']}")
        st.write(f"**Chroma collection:** {info['collection_name']}")
        st.write(f"**Папка с Chroma DB:** `{PERSIST_DIR}`")


st.divider()

# Отрисовка истории чата
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

        if msg.get("sources"):
            render_sources(msg["sources"])

        if msg.get("without_rag"):
            with st.expander("Ответ без RAG для сравнения"):
                st.markdown(msg["without_rag"])


user_question = st.chat_input("Задай вопрос по загруженным документам...")

if user_question:
    st.session_state.messages.append(
        {
            "role": "user",
            "content": user_question,
        }
    )

    with st.chat_message("user"):
        st.markdown(user_question)

    with st.chat_message("assistant"):
        if st.session_state.vectorstore is None:
            answer = "Сначала загрузи и обработай документы, а потом задай вопрос."
            st.warning(answer)

            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": answer,
                }
            )

        else:
            try:
                with st.spinner("Ищу релевантные фрагменты и формирую ответ..."):
                    llm = get_llm(
                        provider=provider,
                        openai_api_key=openai_api_key,
                        openai_chat_model=openai_chat_model,
                        ollama_chat_model=ollama_chat_model,
                        ollama_base_url=ollama_base_url,
                    )

                    answer, sources = answer_with_rag(
                        question=user_question,
                        vectorstore=st.session_state.vectorstore,
                        llm=llm,
                        top_k=top_k,
                    )

                    without_rag_answer = None
                    if compare_without_rag:
                        without_rag_answer = answer_without_rag(user_question, llm)

                st.markdown(answer)
                render_sources(sources)

                if without_rag_answer:
                    with st.expander("Ответ без RAG для сравнения"):
                        st.markdown(without_rag_answer)

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": answer,
                        "sources": sources,
                        "without_rag": without_rag_answer,
                    }
                )

            except Exception as e:
                error_text = f"Ошибка при генерации ответа: {e}"
                st.error(error_text)

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": error_text,
                    }
                )